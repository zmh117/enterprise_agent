from __future__ import annotations

import io
import zipfile

import pytest
from docx import Document
from docx.shared import Inches
from PIL import Image

from app.modules.document_processing.source_validation import (
    normalize_docx_null_image_placeholders,
)
from app.shared.exceptions import NonRetryableExecutionError


_IMAGE_RELATIONSHIP = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"


def _png() -> bytes:
    output = io.BytesIO()
    Image.new("RGB", (8, 8), color="blue").save(output, format="PNG")
    return output.getvalue()


def _docx_with_visible_picture() -> bytes:
    document = Document()
    document.add_paragraph("保留正文")
    document.add_picture(io.BytesIO(_png()), width=Inches(1))
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _add_null_picture_placeholder(
    source: bytes,
    *,
    width: int,
    height: int,
) -> bytes:
    relationship = (
        f'<Relationship Id="rIdNull" Type="{_IMAGE_RELATIONSHIP}" Target="../NULL"/>'
    ).encode()
    drawing = f"""
<w:p><w:r><w:drawing>
  <wp:inline xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">
    <wp:extent cx="{width}" cy="{height}"/>
    <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
      <a:blip xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" r:embed="rIdNull"/>
    </a:graphic>
  </wp:inline>
</w:drawing></w:r></w:p>
""".encode()
    output = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(source)) as archive:
        with zipfile.ZipFile(output, mode="w") as rewritten:
            for info in archive.infolist():
                content = archive.read(info)
                if info.filename == "word/_rels/document.xml.rels":
                    assert b"</Relationships>" in content
                    content = content.replace(
                        b"</Relationships>",
                        relationship + b"</Relationships>",
                    )
                elif info.filename == "word/document.xml":
                    assert b"</w:body>" in content
                    content = content.replace(b"</w:body>", drawing + b"</w:body>")
                rewritten.writestr(info, content)
    return output.getvalue()


def test_normalizer_removes_only_zero_size_null_placeholder_and_keeps_visible_picture() -> None:
    broken = _add_null_picture_placeholder(
        _docx_with_visible_picture(),
        width=635,
        height=0,
    )
    original_snapshot = bytes(broken)

    normalized = normalize_docx_null_image_placeholders(broken, format_code="DOCX")

    assert broken == original_snapshot
    assert normalized != broken
    with zipfile.ZipFile(io.BytesIO(normalized)) as archive:
        assert b"../NULL" not in archive.read("word/_rels/document.xml.rels")
        assert b"rIdNull" not in archive.read("word/document.xml")
        assert "word/media/image1.png" in archive.namelist()
    loaded = Document(io.BytesIO(normalized))
    assert "保留正文" in "\n".join(paragraph.text for paragraph in loaded.paragraphs)
    assert len(loaded.inline_shapes) == 1


def test_normalizer_rejects_visible_null_picture_without_rewriting_source() -> None:
    broken = _add_null_picture_placeholder(
        _docx_with_visible_picture(),
        width=635,
        height=635,
    )
    original_snapshot = bytes(broken)

    with pytest.raises(NonRetryableExecutionError) as error:
        normalize_docx_null_image_placeholders(broken, format_code="DOCX")

    assert error.value.error_code == "docx_null_image_placeholder_unsafe"
    assert broken == original_snapshot


def test_normalizer_returns_exact_source_when_target_flaw_is_absent() -> None:
    source = _docx_with_visible_picture()

    normalized = normalize_docx_null_image_placeholders(source, format_code="DOCX")

    assert normalized is source


def test_normalizer_returns_non_docx_source_without_inspection() -> None:
    source = b"not-an-office-package"

    normalized = normalize_docx_null_image_placeholders(source, format_code="PDF")

    assert normalized is source
